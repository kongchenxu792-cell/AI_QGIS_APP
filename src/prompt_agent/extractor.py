"""文本提取核心 — 支持 .docx 和 .pdf 格式。

优先使用 python-docx / PyMuPDF；缺失时给出明确错误提示。
"""

from __future__ import annotations

import logging
import os
from typing import Callable, Dict

_log = logging.getLogger("prompt_agent.extractor")

# ── 扩展名 → 提取函数 映射表 ──────────────────────────────────────────
_EXTRACTORS: Dict[str, Callable[[str], str]] = {}


def _register(ext: str):
    """装饰器：注册提取函数。"""
    def decorator(fn):
        _EXTRACTORS[ext.lower()] = fn
        return fn
    return decorator


# ── DOCX 提取 ──────────────────────────────────────────────────────────

@_register(".docx")
def extract_docx(file_path: str) -> str:
    """从 .docx 文件提取纯文本（段落 + 表格）。"""
    try:
        from docx import Document  # type: ignore
    except ImportError:
        raise ImportError(
            "缺少 python-docx 库，请执行: pip install python-docx"
        )

    doc = Document(file_path)
    parts: list[str] = []

    # 段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # 表格
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(c for c in cells if c))

    result = "\n".join(parts)
    _log.info("DOCX 提取完成: %s (%d 字符)", file_path, len(result))
    return result


# ── PDF 提取 ───────────────────────────────────────────────────────────

@_register(".pdf")
def extract_pdf(file_path: str) -> str:
    """从 .pdf 文件提取纯文本（逐页），扫描版自动 OCR 兜底。"""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError(
            "缺少 PyMuPDF 库，请执行: pip install PyMuPDF"
        )

    # ── 初始化 Tesseract OCR（用于扫描版 PDF） ──
    _TESSERACT_PATH = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    _ocr_available = os.path.exists(_TESSERACT_PATH)
    if _ocr_available:
        try:
            import pytesseract
            pytesseract.pytesseract.tesseract_cmd = _TESSERACT_PATH
        except ImportError:
            _ocr_available = False
            _log.warning("pytesseract 未安装，扫描版 PDF 将无法提取文本")

    doc = fitz.open(file_path)
    page_count = doc.page_count
    pages_text: list[str] = []
    ocr_page_count = 0

    try:
        for page_num in range(page_count):
            page = doc[page_num]
            text = page.get_text("text")

            if text.strip():
                # 有文本层，直接使用
                pages_text.append(text.strip())
            elif _ocr_available:
                # 扫描页：渲染为图片 → OCR
                pix = page.get_pixmap(dpi=300)
                img_bytes = pix.tobytes("png")
                from PIL import Image
                import io
                img = Image.open(io.BytesIO(img_bytes))
                ocr_result = pytesseract.image_to_string(img, lang="chi_sim+eng")
                if ocr_result.strip():
                    pages_text.append(f"--- 第 {page_num + 1} 页 (OCR) ---\n{ocr_result.strip()}")
                    ocr_page_count += 1
        _log.info(
            "PDF 提取完成: %s (%d 页, %d 字符, OCR %d 页)",
            file_path, page_count, sum(len(t) for t in pages_text), ocr_page_count,
        )
    finally:
        doc.close()

    result = "\n\n".join(pages_text)

    # 全部页面均无文本且 OCR 不可用时给出提示
    if not result.strip():
        if not _ocr_available:
            raise RuntimeError(
                "PDF 为扫描版（无文本层），且 OCR 引擎不可用。"
                "请安装 Tesseract-OCR 并确保 pytesseract 已安装。"
            )
        else:
            raise RuntimeError(
                "OCR 执行完成但未提取到文字，PDF 可能为纯图片且文字不可识别。"
            )

    return result


# ── 统一入口 ───────────────────────────────────────────────────────────

def extract_text(file_path: str) -> str:
    """根据扩展名自动选择提取器，返回纯文本。

    Args:
        file_path: 文档路径（支持 .docx / .pdf）

    Returns:
        提取的纯文本字符串

    Raises:
        ValueError: 不支持的文件格式
        FileNotFoundError: 文件不存在
        ImportError: 缺少必要的第三方库
    """
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()
    extractor = _EXTRACTORS.get(ext)

    if extractor is None:
        supported = ", ".join(_EXTRACTORS.keys())
        raise ValueError(
            f"不支持的文件格式 '{ext}'，当前支持: {supported}"
        )

    return extractor(file_path)